"""MCP Server for python-docx — create and manipulate Word documents."""

from __future__ import annotations

import asyncio
import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from mcp.server import Server
from mcp.server.stdio import stdio_server
from mcp.types import Tool, TextContent

logger = logging.getLogger(__name__)

server = Server("docx-server")

# ── In-memory state ──────────────────────────────────────────────
_doc: Document | None = None


def _ensure_doc() -> Document:
    if _doc is None:
        raise ValueError("No document created yet. Call create_document first.")
    return _doc


_ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


# ── Tools ────────────────────────────────────────────────────────

@server.list_tools()
async def list_tools() -> list[Tool]:
    return [
        Tool(
            name="create_document",
            description="Create a new blank Word document in memory",
            inputSchema={"type": "object", "properties": {}, "required": []},
        ),
        Tool(
            name="add_heading",
            description="Add a heading to the document",
            inputSchema={
                "type": "object",
                "properties": {
                    "text": {"type": "string", "description": "Heading text"},
                    "level": {
                        "type": "integer",
                        "description": "Heading level (1-9, where 1 is the largest)",
                        "default": 1,
                    },
                },
                "required": ["text"],
            },
        ),
        Tool(
            name="add_paragraph",
            description="Add a paragraph to the document with optional formatting",
            inputSchema={
                "type": "object",
                "properties": {
                    "text": {"type": "string", "description": "Paragraph text"},
                    "bold": {"type": "boolean", "description": "Make text bold"},
                    "italic": {"type": "boolean", "description": "Make text italic"},
                    "underline": {"type": "boolean", "description": "Underline text"},
                    "alignment": {
                        "type": "string",
                        "enum": ["left", "center", "right", "justify"],
                        "description": "Text alignment",
                    },
                    "font_size": {"type": "integer", "description": "Font size in points (e.g. 12)"},
                },
                "required": ["text"],
            },
        ),
        Tool(
            name="add_table",
            description="Add a table from a 2D array of data; first row is treated as header",
            inputSchema={
                "type": "object",
                "properties": {
                    "rows": {
                        "type": "array",
                        "items": {"type": "array", "items": {"type": "string"}},
                        "description": "2D array of cell values (first row becomes the header)",
                    },
                    "style": {
                        "type": "string",
                        "description": "Table style name (e.g. 'Table Grid', 'Light Shading Accent 1')",
                        "default": "Table Grid",
                    },
                },
                "required": ["rows"],
            },
        ),
        Tool(
            name="add_picture",
            description="Embed an image into the document",
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Absolute path to the image file"},
                    "width_inches": {"type": "number", "description": "Width in inches (optional)"},
                    "height_inches": {"type": "number", "description": "Height in inches (optional)"},
                },
                "required": ["path"],
            },
        ),
        Tool(
            name="add_page_break",
            description="Insert a page break at the current cursor position",
            inputSchema={"type": "object", "properties": {}, "required": []},
        ),
        Tool(
            name="get_document_info",
            description="Get info about the current document: paragraph count, sections, tables",
            inputSchema={"type": "object", "properties": {}, "required": []},
        ),
        Tool(
            name="save_document",
            description="Save the document to a .docx file on disk",
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Absolute path for the output .docx file"},
                },
                "required": ["path"],
            },
        ),
    ]


@server.call_tool()
async def call_tool(name: str, arguments: dict) -> list[TextContent]:
    global _doc

    try:
        if name == "create_document":
            _doc = Document()
            return [TextContent(type="text", text="Blank document created.")]

        elif name == "add_heading":
            doc = _ensure_doc()
            level = max(1, min(9, arguments.get("level", 1)))
            doc.add_heading(arguments["text"], level=level)
            return [TextContent(type="text", text=f"Added heading (level {level}): {arguments['text']}")]

        elif name == "add_paragraph":
            doc = _ensure_doc()
            para = doc.add_paragraph()
            run = para.add_run(arguments["text"])
            if arguments.get("bold"):
                run.bold = True
            if arguments.get("italic"):
                run.italic = True
            if arguments.get("underline"):
                run.underline = True
            if arguments.get("font_size"):
                run.font.size = Pt(arguments["font_size"])
            if arguments.get("alignment"):
                para.alignment = _ALIGNMENT_MAP.get(arguments["alignment"], WD_ALIGN_PARAGRAPH.LEFT)
            return [TextContent(type="text", text=f"Added paragraph: {arguments['text'][:80]}")]

        elif name == "add_table":
            doc = _ensure_doc()
            rows = arguments["rows"]
            if not rows or not rows[0]:
                raise ValueError("rows must be a non-empty 2D array")
            ncols = len(rows[0])
            table = doc.add_table(rows=len(rows), cols=ncols)
            style_name = arguments.get("style", "Table Grid")
            try:
                table.style = style_name
            except KeyError:
                pass  # style not found, use default
            for i, row_data in enumerate(rows):
                for j, cell_text in enumerate(row_data):
                    if j < ncols:
                        cell = table.cell(i, j)
                        cell.text = str(cell_text)
                        if i == 0:  # header row
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.bold = True
            return [TextContent(type="text", text=f"Added table with {len(rows)} rows × {ncols} cols.")]

        elif name == "add_picture":
            doc = _ensure_doc()
            path = arguments["path"]
            if not Path(path).is_file():
                return [TextContent(type="text", text=f"Error: file not found — {path}")]
            kwargs = {}
            if arguments.get("width_inches"):
                kwargs["width"] = Inches(arguments["width_inches"])
            if arguments.get("height_inches"):
                kwargs["height"] = Inches(arguments["height_inches"])
            doc.add_picture(path, **kwargs)
            return [TextContent(type="text", text=f"Added picture: {path}")]

        elif name == "add_page_break":
            doc = _ensure_doc()
            doc.add_page_break()
            return [TextContent(type="text", text="Page break added.")]

        elif name == "get_document_info":
            doc = _ensure_doc()
            info_lines = [
                f"Paragraphs: {len(doc.paragraphs)}",
                f"Sections: {len(doc.sections)}",
                f"Tables: {len(doc.tables)}",
                f"Inline shapes (images): {len(doc.inline_shapes)}",
            ]
            return [TextContent(type="text", text="\n".join(info_lines))]

        elif name == "save_document":
            doc = _ensure_doc()
            path = arguments["path"]
            out = Path(path)
            out.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(out))
            size_kb = out.stat().st_size / 1024
            return [TextContent(type="text", text=f"Document saved: {path} ({size_kb:.1f} KB)")]

        else:
            return [TextContent(type="text", text=f"Unknown tool: {name}")]

    except ValueError as e:
        return [TextContent(type="text", text=str(e))]
    except Exception as e:
        logger.exception("Tool error")
        return [TextContent(type="text", text=f"Error: {e}")]


# ── Entry point ──────────────────────────────────────────────────

def main():
    logger.info("Starting docx-server...")
    asyncio.run(_run())


async def _run():
    async with stdio_server() as (read_stream, write_stream):
        await server.run(read_stream, write_stream, server.create_initialization_options())


if __name__ == "__main__":
    main()
